import streamlit as st
import google.generativeai as genai
import os
import logging
from datetime import datetime
import urllib.parse
from fpdf import FPDF
from docx import Document
import tempfile
from functools import lru_cache
import time

# Initialize logger with a NullHandler by default
logger = logging.getLogger(__name__)
logger.addHandler(logging.NullHandler())

# Only set up file logging if ENABLE_LOGGING environment variable is set to "true"
if os.getenv("ENABLE_LOGGING", "false").lower() == "true":
    # Create logs directory if it doesn't exist
    logs_dir = "logs"
    if not os.path.exists(logs_dir):
        os.makedirs(logs_dir)

    # Set up logging
    log_filename = os.path.join(logs_dir, f'paper_generator_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log')
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_filename),
            logging.StreamHandler()  # This will also print to console
        ]
    )
    logger = logging.getLogger(__name__)
    
    # Log startup information
    logger.info("Starting Paper Generator Application")
    logger.info(f"Logging to: {log_filename}")

# Configure Gemini API
genai.configure(api_key="AIzaSyB_cNSUmfNKahsm39hr7xPbXxNCta-GC50")

# Initialize the model with safety settings
generation_config = {
    "temperature": 0.7,
    "top_p": 0.8,
    "top_k": 40,
    "max_output_tokens": 1024,
}

safety_settings = [
    {
        "category": "HARM_CATEGORY_HARASSMENT",
        "threshold": "BLOCK_NONE"
    },
    {
        "category": "HARM_CATEGORY_HATE_SPEECH",
        "threshold": "BLOCK_NONE"
    },
    {
        "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
        "threshold": "BLOCK_NONE"
    },
    {
        "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
        "threshold": "BLOCK_NONE"
    },
]

model = genai.GenerativeModel(
    model_name='gemini-pro',
    generation_config=generation_config,
    safety_settings=safety_settings
)

def clean_ai_response(response_text: str) -> str:
    """Clean up AI response text and handle markdown formatting"""
    logger.debug(f"Cleaning response text: {response_text}")
    
    try:
        # Remove markdown code blocks if present
        cleaned_text = response_text.strip()
        if cleaned_text.startswith("```json"):
            cleaned_text = cleaned_text.replace("```json", "").replace("```", "")
        elif cleaned_text.startswith("```"):
            cleaned_text = cleaned_text.replace("```", "")
        cleaned_text = cleaned_text.strip()
        
        logger.debug(f"Cleaned text: {cleaned_text}")
        return cleaned_text
    except Exception as e:
        logger.error(f"Error cleaning response text: {e}")
        raise

def generate_text(prompt: str, temperature: float) -> str:
    """Generate text using the AI model"""
    try:
        response = model.generate_content(prompt)
        return clean_ai_response(response.text)
    except Exception as e:
        logger.error(f"Error generating text: {e}")
        st.error("Failed to generate content")
        return ""

def generate_single_paper(subject: str) -> str:
    """Generate a single paper for a given subject"""
    try:
        rate_limiter.wait_if_needed()  # Rate limit API calls
        
        # Construct prompts for each section
        mcq_prompt = (
            f"Generate {template['mcq']} multiple-choice questions for Class {class_selected}, "
            f"Subject: {subject}. Difficulty: {difficulty}"
        )
        fill_blanks_prompt = (
            f"Generate {template['fill']} fill-in-the-blank questions for Class {class_selected}, "
            f"Subject: {subject}. Difficulty: {difficulty}"
        )
        small_questions_prompt = (
            f"Generate {template['small']} small answer questions for Class {class_selected}, "
            f"Subject: {subject}. Difficulty: {difficulty}"
        )
        big_questions_prompt = (
            f"Generate {template['big']} long answer questions for Class {class_selected}, "
            f"Subject: {subject}. Difficulty: {difficulty}"
        )

        # Generate sections using cached function
        mcqs = generate_text_cached(mcq_prompt, temperature=0.7)
        fill_blanks = generate_text_cached(fill_blanks_prompt, temperature=0.7)
        small_questions = generate_text_cached(small_questions_prompt, temperature=0.7)
        big_questions = generate_text_cached(big_questions_prompt, temperature=0.7)

        # Combine sections
        paper = (
            f"Subject: {subject}\n\n"
            f"MCQs:\n{mcqs}\n\n"
            f"Fill in the Blanks:\n{fill_blanks}\n\n"
            f"Small Questions:\n{small_questions}\n\n"
            f"Big Questions:\n{big_questions}"
        )

        # Review and regenerate
        review_comments = review_paper(paper)
        reviewed_paper = regenerate_based_on_review(paper, review_comments)
        return reviewed_paper

    except Exception as e:
        logger.error(f"Error generating paper for {subject}: {e}")
        raise PaperGenerationError(f"Failed to generate paper for {subject}")

def review_paper(paper):
    # Simulate querying an LLM for paper review
    review_prompt = f"Review the following test paper and provide suggestions for improvement:\n\n{paper}"
    review_comments = generate_text(review_prompt, temperature=0.5)
    return review_comments

def regenerate_based_on_review(original_paper, review_comments):
    # Simulate querying an LLM to regenerate the paper based on review comments
    regeneration_prompt = f"Improve the following test paper based on the review comments:\n\nOriginal Paper:\n{original_paper}\n\nReview Comments:\n{review_comments}"
    regenerated_paper = generate_text(regeneration_prompt, temperature=0.7)
    return regenerated_paper

def generate_answer_key(paper_content):
    """Generate answer key for the test paper"""
    prompt = f"""
    Create an answer key for this test paper:
    {paper_content}
    
    Format answers with explanations for each question.
    Return in this format:
    MCQs:
    1. Answer (explanation)
    2. Answer (explanation)
    
    Fill in the Blanks:
    1. Answer (explanation)
    
    Short Answers:
    1. Key points to look for in answer
    
    Long Answers:
    1. Detailed evaluation criteria and sample answer
    """
    return generate_text(prompt, temperature=0.3)

def add_difficulty_distribution():
    st.sidebar.subheader("Question Difficulty Distribution")
    easy = st.sidebar.slider("Easy Questions (%)", 0, 100, 30)
    medium = st.sidebar.slider("Medium Questions (%)", 0, 100, 40)
    hard = st.sidebar.slider("Hard Questions (%)", 0, 100, 30)
    
    total = easy + medium + hard
    if total != 100:
        return None
    return {"easy": easy, "medium": medium, "hard": hard}

def calculate_time_duration(num_mcq, num_fill_blanks, num_small_qs, num_big_qs):
    """Calculate recommended time duration for the test"""
    mcq_time = num_mcq * 1  # 1 minute per MCQ
    fill_time = num_fill_blanks * 1  # 1 minute per fill in blank
    small_qs_time = num_small_qs * 5  # 5 minutes per small question
    big_qs_time = num_big_qs * 15  # 15 minutes per big question
    
    total_time = mcq_time + fill_time + small_qs_time + big_qs_time
    return total_time

def generate_marking_scheme(num_mcq, num_fill_blanks, num_small_qs, num_big_qs):
    total_marks = 0
    marking_scheme = {
        "MCQs": {"per_question": 1, "total": num_mcq * 1},
        "Fill_Blanks": {"per_question": 1, "total": num_fill_blanks * 1},
        "Small_Questions": {"per_question": 3, "total": num_small_qs * 3},
        "Big_Questions": {"per_question": 5, "total": num_big_qs * 5}
    }
    total_marks = sum(section["total"] for section in marking_scheme.values())
    return marking_scheme, total_marks

def add_paper_templates():
    templates = {
        "Standard": {"mcq": 10, "fill": 5, "small": 5, "big": 3},
        "MCQ Focus": {"mcq": 20, "fill": 5, "small": 3, "big": 1},
        "Descriptive": {"mcq": 5, "fill": 5, "small": 8, "big": 5},
        "Quick Test": {"mcq": 5, "fill": 5, "small": 3, "big": 1}
    }
    
    selected_template = st.sidebar.selectbox(
        "Choose Paper Template",
        list(templates.keys())
    )
    return templates[selected_template]

def export_paper(paper_content, format_type):
    if format_type == "PDF":
        return create_pdf(paper_content)
    elif format_type == "Word":
        return create_word_doc(paper_content)
    elif format_type == "Plain Text":
        return create_text_file(paper_content)

def create_pdf(paper_content: str) -> bytes:
    """Create a PDF document from the paper content"""
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Split content into lines and add to PDF
        lines = paper_content.split('\n')
        for line in lines:
            # Encode to handle special characters
            encoded_line = line.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, encoded_line)
            
        # Save to bytes buffer
        return pdf.output(dest='S').encode('latin-1')
    except Exception as e:
        logger.error(f"Error creating PDF: {e}")
        st.error("Failed to create PDF document")
        return None

def create_word_doc(paper_content: str) -> bytes:
    """Create a Word document from the paper content"""
    try:
        doc = Document()
        doc.add_heading('Test Paper', 0)
        
        # Add content paragraphs
        for paragraph in paper_content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph)
        
        # Save to temporary file and read bytes
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc.save(tmp_file.name)
            with open(tmp_file.name, 'rb') as file:
                return file.read()
    except Exception as e:
        logger.error(f"Error creating Word document: {e}")
        st.error("Failed to create Word document")
        return None

def create_text_file(paper_content: str) -> bytes:
    """Create a plain text file from the paper content"""
    try:
        return paper_content.encode('utf-8')
    except Exception as e:
        logger.error(f"Error creating text file: {e}")
        st.error("Failed to create text file")
        return None

def generate_papers_with_progress() -> str:
    """Generate papers with progress tracking"""
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    try:
        generated_papers = []
        total_subjects = len(subjects_selected)
        
        for idx, subject in enumerate(subjects_selected):
            status_text.text(f"Generating {subject} paper...")
            paper = generate_single_paper(subject)
            generated_papers.append(paper)
            progress_bar.progress((idx + 1) / total_subjects)
            
        status_text.text("Generation complete!")
        return "\n\n---\n\n".join(generated_papers)
    except Exception as e:
        status_text.text("Generation failed!")
        raise e

def create_test_paper():
    st.title("CBSE Test Paper Generator")
    
    # Initialize session state if not exists
    if 'generated_papers' not in st.session_state:
        st.session_state.generated_papers = None
    if 'downloads' not in st.session_state:
        st.session_state.downloads = []
    
    # Initialize these as global variables
    global template, difficulty_dist, class_selected, subjects_selected, difficulty
    
    # Add templates
    template = add_paper_templates()
    
    # Add difficulty distribution
    difficulty_dist = add_difficulty_distribution()
    if not difficulty_dist:
        st.error("Please adjust difficulty distribution to total 100%")
        return
    
    # Define default subjects
    default_subjects = [
        "Maths", 
        "Science (Physics)", "Science (Biology)", "Science (Chemistry)",
        "English (Literature)", "English (Grammar)",
        "Hindi", "Kannada",
        "Social Science (History)", "Social Science (Geography)",
        "Social Science (Political Science)", "Social Science (Economics)"
    ]
    
    # Original inputs
    class_selected = st.selectbox("Select Class:", ["4th", "5th", "6th", "7th", "8th", "9th", "10th"])
    subjects_selected = st.multiselect("Select Subject(s):", options=default_subjects)
    
    # Difficulty level
    difficulty = st.radio("Select Difficulty Level:", ["Easy", "Moderate", "Difficult"])
    
    if not subjects_selected:
        st.warning("Please select at least one subject")
        return
    
    # Select format and number of questions
    st.markdown("### Customize Test Format")
    
    # Use template values instead of separate number inputs
    num_mcq = template["mcq"]
    num_fill_blanks = template["fill"]
    num_small_qs = template["small"]
    num_big_qs = template["big"]
    
    # Optional custom format
    custom_format = st.text_area(
        "Provide Existing Format (Optional):",
        placeholder="Paste your existing test format here..."
    )
    
    # Calculate time and marks
    time_duration = calculate_time_duration(num_mcq, num_fill_blanks, 
                                         num_small_qs, num_big_qs)
    marking_scheme, total_marks = generate_marking_scheme(num_mcq, 
                                                        num_fill_blanks,
                                                        num_small_qs, 
                                                        num_big_qs)
    
    # Display paper details
    st.info(f"Estimated Time: {time_duration} minutes")
    st.info(f"Total Marks: {total_marks}")
    
    # Add option to generate answer key
    generate_answers = st.checkbox("Generate Answer Key", value=True, 
                                 help="Uncheck to skip generating answers")
    
    if st.button("Generate Test Papers", key="generate_btn"):
        with st.spinner("Generating papers..."):
            generated_papers = generate_papers_with_progress()
            if generated_papers:
                # Store in session state
                st.session_state.generated_papers = generated_papers
                st.session_state.show_export = True  # Flag to show export section
                
                # Display papers
                st.subheader("Generated Test Papers")
                st.text_area("Test Papers:", generated_papers, height=500)
                
                # Generate answers if selected
                if generate_answers:
                    with st.spinner("Generating answer key..."):
                        answer_keys = generate_answer_key(generated_papers)
                        st.session_state.answer_keys = answer_keys  # Store in session state
                        st.subheader("Answer Key")
                        st.text_area("Answers:", answer_keys, height=300)

    # Show export section if papers were generated
    if hasattr(st.session_state, 'show_export') and st.session_state.show_export:
        st.subheader("Export Options")
        
        col1, col2 = st.columns([3, 1])
        
        with col1:
            format_type = st.selectbox(
                "Select Format",
                ["PDF", "Word", "Plain Text"],
                key="export_format"
            )
        
        with col2:
            if st.button("Export", key="export_btn", use_container_width=True):
                try:
                    with st.spinner("Preparing download..."):
                        # Get content from session state
                        export_content = st.session_state.generated_papers
                        if generate_answers and hasattr(st.session_state, 'answer_keys'):
                            export_content += "\n\n=== ANSWER KEY ===\n\n" + st.session_state.answer_keys
                        
                        exported_file = export_paper(export_content, format_type)
                        
                        if exported_file:
                            file_extension = {
                                "PDF": ".pdf",
                                "Word": ".docx",
                                "Plain Text": ".txt"
                            }[format_type]
                            
                            mime_type = {
                                "PDF": "application/pdf",
                                "Word": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                "Plain Text": "text/plain"
                            }[format_type]
                            
                            # Create unique filename
                            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                            filename = f"test_paper_{timestamp}{file_extension}"
                            
                            # Add to downloads history
                            if 'downloads' not in st.session_state:
                                st.session_state.downloads = []
                            st.session_state.downloads.append({
                                'filename': filename,
                                'timestamp': datetime.now()
                            })
                            
                            # Create container for download button
                            download_container = st.container()
                            with download_container:
                                st.download_button(
                                    label="📥 Download File",
                                    data=exported_file,
                                    file_name=filename,
                                    mime=mime_type,
                                    key=f"download_{timestamp}"
                                )
                                st.success(f"✅ Test paper exported as {format_type}!")
                            
                except Exception as e:
                    logger.error(f"Export failed: {e}")
                    st.error(f"Failed to export: {str(e)}")

# Add SessionState class to better manage state
class SessionState:
    def __init__(self):
        self.template = None
        self.difficulty_dist = None
        self.class_selected = None
        self.subjects_selected = None
        self.difficulty = None
        self.generated_papers = None

def initialize_session_state():
    if 'state' not in st.session_state:
        st.session_state.state = SessionState()

# Add configuration class
class Config:
    SUBJECTS = {
        "Sciences": ["Physics", "Biology", "Chemistry"],
        "Languages": ["English Literature", "English Grammar", "Hindi", "Kannada"],
        "Social Sciences": ["History", "Geography", "Political Science", "Economics"],
        "Mathematics": ["Maths"]
    }
    
    TEMPLATES = {
        "Standard": {"mcq": 10, "fill": 5, "small": 5, "big": 3},
        "MCQ Focus": {"mcq": 20, "fill": 5, "small": 3, "big": 1},
        "Descriptive": {"mcq": 5, "fill": 5, "small": 8, "big": 5},
        "Quick Test": {"mcq": 5, "fill": 5, "small": 3, "big": 1}
    }

@st.cache_data(ttl=3600)  # Cache for 1 hour
def generate_text_cached(prompt: str, temperature: float) -> str:
    """Cached version of generate_text to improve performance"""
    return generate_text(prompt, temperature)

class PaperGenerationError(Exception):
    """Custom exception for paper generation errors"""
    pass

def handle_generation_error(func):
    """Decorator for handling generation errors"""
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            logger.error(f"Error in {func.__name__}: {str(e)}")
            st.error(f"Failed to generate paper: {str(e)}")
            return None
    return wrapper

class RateLimiter:
    def __init__(self, calls_per_second=2):
        self.calls_per_second = calls_per_second
        self.last_call = 0

    def wait_if_needed(self):
        elapsed = time.time() - self.last_call
        if elapsed < 1/self.calls_per_second:
            time.sleep(1/self.calls_per_second - elapsed)
        self.last_call = time.time()

rate_limiter = RateLimiter()

def track_downloads():
    """Track paper downloads"""
    if 'downloads' not in st.session_state:
        st.session_state.downloads = []
    
    if st.session_state.downloads:  # Only show if there are downloads
        st.sidebar.markdown("### Recent Downloads")
        for download in st.session_state.downloads[-5:]:  # Show last 5 downloads
            filename = download.get('filename', 'Unnamed file')
            st.sidebar.text(f"📄 {filename}")

# Remove duplicate code at the bottom and just call create_test_paper()
create_test_paper()

# Footer
st.markdown("---")
st.markdown("*CBSE Test Paper Generator powered by AI*")
